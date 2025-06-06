import os
import pandas as pd
import datetime
import docx
from xhtml2pdf import pisa


# ___UTILS___
# treat empty datetime
def to_string_custom(date_in):
    if date_in:
        date_string = date_in.strftime('%d-%m-%Y')
    else:
        date_string = "Non renseignée"
    return date_string


# ___EXPORT___

# export .csv from a dataframe and some params
# key needed for param object
# - ["nomenclature_ref"]
# - ["nomenclature_name"]
# - ["docs"]
# - ["redacteur"]
# - ["date_validation"]
# - ["date_expiration"]
# - ["levels"]
def export_csv_nomenclature(params_in, df_nomenclature_in, folder_output):
    df_export = df_nomenclature_in.copy()
    df_export.rename(columns={"Code": "code", "Description": "description", "Commentaire": "commentaire"}, inplace=True)
    for level in range(1, params_in["levels"] + 1):
        df_export.rename(columns={"Libellé niveau " + str(level): "label_n" + str(level)}, inplace=True)
    file_out = params_in["nomenclature_id"] + ".csv"
    # create csv dir if not exist
    try:
        os.mkdir(os.path.join(folder_output, "csv"))
        print("Creating folder " + str(os.path.join(folder_output, "csv")) + " ...")
    except FileExistsError:
        pass
    # export
    df_export.to_csv(os.path.join(folder_output, "csv", file_out))
    return


def export_pdf_nomenclature(params_in, df_nomenclature_in, folder_output):
    df_export = df_nomenclature_in.copy()
    html_start = f"""
        <head>
        <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
        """
    html_style = """
    <style>.dataframe 
    th{
        background: linear-gradient(71deg, rgb(20, 0, 94) 0%, rgb(25, 18, 126) 100%);
        padding: 10px;
        font-family: Arial;
        color: #ffffff;
        border:1px;
        text-align:left !important;
        margin-top:1px;
        margin-bottom:1px;
        padding-top:1px;
        padding-bottom:1px;
        }
    td{
        border:1px;
        font-family: Arial;
        margin-top:1px;
        margin-bottom:1px;
        margin-left:2px;
        padding-top:2px;
        padding-bottom:1px;
    }
    </style>
    </head>
    <body>
    """
    html_end = """</body>"""
    html_corps = df_export.to_html(index=False, na_rep="")

    # concatenate html parts
    html = html_start + html_style + html_corps + html_end

    file_out = params_in["nomenclature_id"] + ".pdf"
    # create pdf dir if not exist
    try:
        os.mkdir(os.path.join(folder_output, "pdf"))
        print("Creating folder " + str(os.path.join(folder_output, "pdf")) + " ...")
    except FileExistsError:
        pass
    # export
    # open output file for writing (truncated binary)
    pdf_file = open(os.path.join(folder_output, "pdf", file_out), "w+b")
    # convert HTML to PDF
    pisa_status = pisa.CreatePDF(
        html,  # the HTML to convert
        dest=pdf_file)  # file handle to recieve result
    # close output file
    pdf_file.close()
    return


# transform df dataframe to a doc objet
def df_nomenclature_to_doc(params_in, df_nomenclature_in, doc=None, style='Medium Shading 1 Accent 1'):
    if doc is None:
        doc = docx.Document()

    # Add header
    doc.add_heading(params_in["nomenclature_ref"] + "-" + params_in["nomenclature_name"], level=1)
    doc.add_paragraph()

    # Add table
    table_date = doc.add_table(rows=2, cols=2, style=style)

    # table data
    table_date.rows[0].cells[0].text = "Date de validation"
    table_date.rows[1].cells[0].text = to_string_custom(params_in["date_validation"] if "date_validation" in params_in else "")

    table_date.rows[0].cells[1].text = "Date d'expiration"
    table_date.rows[1].cells[1].text = to_string_custom(params_in["date_expiration"] if "date_expiration" in params_in else "")

    # table style
    table_date.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table_date.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Rédacteur(s) : " + params_in["redacteur"] if "redacteur" in params_in else "")
    doc.add_paragraph("Description : " + params_in["nomenclature_description"] if "nomenclature_description" in params_in else "")

    # Add paragraph
    # doc.add_paragraph('This table represents the fields and types defined in the JSON schema.')

    n_cols_tot = df_nomenclature_in.shape[1]
    # extra row is so we can add the header row
    table = doc.add_table(df_nomenclature_in.shape[0] + 1, cols=n_cols_tot, style=style)

    # add the header rows.
    for j in range(df_nomenclature_in.shape[-1]):
        table.cell(0, j).text = df_nomenclature_in.columns[j]

    # add the rest of the data frame
    for i in range(df_nomenclature_in.shape[0]):
        for j in range(df_nomenclature_in.shape[-1]):
            table.cell(i + 1, j).text = str(df_nomenclature_in.fillna("").values[i, j])

    # Specify the column widths (has to be on cells for Word) | Ref.: https://stackoverflow.com/a/43053996
    table.autofit = False
    table.allow_autofit = False
    # column_widths = (3, 3.5, 2, 2.5, 8, 3)  # Widths in centimeters
    # set_col_widths(table, column_widths)

    return doc


# export .pdf and .docx
def export_docx_nomenclature(params_in, df_nomenclature_in, folder_output):
    # generate docx
    doc_export = df_nomenclature_to_doc(params_in, df_nomenclature_in)

    # .docx
    # create word dir if not exist
    file_out_docx = params_in["nomenclature_id"] + ".docx"
    try:
        os.mkdir(os.path.join(folder_output, "word"))
        print("Creating folder " + str(os.path.join(folder_output, "word")) + " ...")
    except FileExistsError:
        pass
    # export
    doc_export.save(os.path.join(folder_output, "word", file_out_docx))
    return


# transform sommaire dataframe to a doc
def df_sommaire_to_doc(df_sommaire_in, doc=None, style='Medium Shading 1 Accent 1'):
    if doc is None:
        doc = docx.Document()

    # Add header
    doc.add_heading("Nomenclature Hub Santé", level=1)
    doc.add_paragraph()

    df_sommaire_in.rename(columns={"nomenclature_name": "Nomenclature",
                                   "file_name": "Nom de fichier",
                                   "referentiel": "Référentiel d'origine de la norme"}, inplace=True)

    table = doc.add_table(df_sommaire_in.shape[0] + 1, cols=df_sommaire_in.shape[1], style=style)

    # add the header rows.
    for j in range(df_sommaire_in.shape[-1]):
        table.cell(0, j).text = df_sommaire_in.columns[j]

    # add the rest of the data frame
    for i in range(df_sommaire_in.shape[0]):
        for j in range(df_sommaire_in.shape[-1]):
            table.cell(i + 1, j).text = str(df_sommaire_in.fillna("").values[i, j])

    return doc


def export_sommaire(df_sommaire_in, folder_output):
    # generate docx
    doc_export = df_sommaire_to_doc(df_sommaire_in)

    # .docx
    file_out_docx = "Sommaire Nomenclature Hub" + ".docx"

    # export
    doc_export.save(os.path.join(folder_output, file_out_docx))
    return


# ___PARSING SHEET___
def parse_sheet(filename_in, sheet_name_in):
    full_sheet = pd.read_excel(filename_in, sheet_name=sheet_name_in)
    # reading rows till the beginning of the nomenclature table (which begins with "Code")
    # first we find the row index of the first row containing "Code"
    # then we read the rows from the beginning of the file to this row
    beginning_row = full_sheet[full_sheet.iloc[:, 0] == "Code"].index[0]

    df_params = pd.read_excel(filename_in, sheet_name=sheet_name_in,
                              nrows=beginning_row+1, header=None, index_col=0,
                              dtype={"levels": int, "date_validation": datetime, "date_expiration": datetime},
                              keep_default_na=False, na_values=['NULL', 'null', 'NaN', 'nan', 'None', 'none', ''])
    df_params.fillna("", inplace=True)
    params = {}
    # If df_params has "Identifiant nomenclature", we use it as "nomenclature_id" parameter, otherwise we use the sheet name
    params["nomenclature_ref"] = df_params.loc["Référentiel nomenclature"].iloc[0]
    params["nomenclature_name"] = df_params.loc["Nom Nomenclature"].iloc[0]
    params["nomenclature_id"] = df_params.loc["Identifiant nomenclature"].iloc[0] if "Identifiant nomenclature" in df_params.index else params["nomenclature_ref"] + "-" + params["nomenclature_name"]
    params["levels"] = df_params.loc["Nombre de niveau"].iloc[0]
    # parsing levels as int
    try:
        params["levels"] = int(params["levels"])
    except ValueError:
        print("Field 'Nombre de niveau' shall be an integer")
        raise
    params["nomenclature_description"] = df_params.loc["Description"].iloc[0]
    # print(params)
    # skipping till beginning of actual nomenclature table
    df_nomenclature = pd.read_excel(filename_in, sheet_name=sheet_name_in, skiprows=beginning_row+1,
                                    keep_default_na=False, na_values=['NULL', 'null', 'NaN', 'nan', 'None', 'none', ''])
    # checking levels
    for level in range(1, params["levels"] + 1):
        if ("Libellé niveau " + str(level)) not in df_nomenclature.columns:
            print("Colonne de niveau manquante ou libellé de niveau de colonne incorrect")
            raise ValueError
    # checking columns name minimal
    if not all(item in df_nomenclature.columns for item in ["Code", "Description", "Commentaire"]):
        print("Libellé de colonne manquant")
        raise ValueError
    # return
    return params, df_nomenclature


# ___PARSING A FILE THEN A FOLDER___
def parse_excel(filename_in, df_sommaire_in, folder_output):
    sheet_names = pd.ExcelFile(filename_in).sheet_names
    # iterate over sheets
    for sheet_i in sheet_names:
        if sheet_i.startswith("#"):
            print("Skipping " + filename_in + " : sheet " + sheet_i + " ...")
            continue  # we end this loop and go to the next sheet
        print("Processing " + filename_in + " : sheet " + sheet_i + " ...")
        params_i, df_nomenclature_i = parse_sheet(filename_in, sheet_i)
        nom_fichier_i = params_i["nomenclature_ref"] + "-" + params_i["nomenclature_name"]

        # complete sommaire
        df_i = pd.DataFrame({
                            "nomenclature_id": params_i["nomenclature_id"],
                            "nomenclature_name": params_i["nomenclature_name"],
                             "file_name": nom_fichier_i,
                             "referentiel": params_i["nomenclature_ref"]},
                            index=[0])
        df_sommaire_in = pd.concat([df_i, df_sommaire_in.loc[:]]).reset_index(drop=True)

        # generate export
        export_csv_nomenclature(params_i, df_nomenclature_i, folder_output)
        export_pdf_nomenclature(params_i, df_nomenclature_i, folder_output)
        export_docx_nomenclature(params_i, df_nomenclature_i, folder_output)
    return df_sommaire_in


def parse_folder(folder_in, folder_output):
    # initialize empty df for summary
    df_sommaire = pd.DataFrame(data={},
                               columns=["nomenclature_id",
                                        "nomenclature_name",
                                        "file_name",
                                        "referentiel"],
                               dtype=str)

    for filename in os.listdir(folder_in):
        # check extension 
        if filename.endswith('.xlsx'):
            print("Processing " + filename + " ...")
            # processing
            df_sommaire = parse_excel(os.path.join(folder_in, filename), df_sommaire, folder_output)
        else:
            print(filename + " is not a valid nomenclature file.")
    print("Processing summary ...")
    export_sommaire(df_sommaire, folder_output)
    print("Job done.")
    return
