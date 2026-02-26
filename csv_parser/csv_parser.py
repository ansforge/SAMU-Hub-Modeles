import copy
import re
from enum import StrEnum
import pandas as pd
import json
from jsonpath_ng import parse
import yaml
import docx
import argparse
from datetime import date
import warnings
import uml_generator
import os

from pathlib import Path
# Improving panda printing | Ref.: https://stackoverflow.com/a/11711637
pd.set_option('display.max_rows', 500)
pd.set_option('display.max_columns', 500)
pd.set_option('display.width', 1000)

# Ignoring Openpyxl Excel's warnings | Ref.: https://stackoverflow.com/a/64420416
warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

all_model_types = []
first_codeandlabel_name = ""
first_codeandlabel_properties = []


def run(sheet, name, version, perimeter_filter, model_type, filepath):
    class Color:
        ORANGE = '\033[93m'
        RED = '\033[91m'
        PURPLE = '\033[95m'
        BOLD = '\033[1m'
        UNDERLINE = '\033[4m'
        END = '\033[0m'

    version = version or date.today().strftime("%y.%m.%d")
    print(f"{Color.BOLD}{Color.UNDERLINE}{Color.PURPLE}"
          f"Building version {version} of {perimeter_filter} {sheet} sheet "
          f"into {name} schema..."
          f"{Color.END}")

    RUN_DOCX_OUTPUT_EXAMPLES = False

    DATA_DEPTH = 6  # nombre de niveaux de données
    HEADER_LINE = 7  # ligne avec les en-têtes (ID Excel - 1)

    class FormatFlags(StrEnum):
        NOMENCLATURE = 'NOMENCLATURE'
        ENUM = 'ENUM'
        REGEX = 'REGEX'


    def get_params_from_sheet(sheet):
        """ Automatically get the number of rows and columns to use for this sheet. """
        full_df = pd.read_excel(filepath, sheet_name=sheet, header=None)
        # Getting modelName from cell A1
        modelName = full_df.iloc[0, 0]
        # Get line {HEADER_LINE - 1} to find which columns have 'Périmètre' flag
        perimeter_row = full_df.iloc[HEADER_LINE - 1, :]
        # Save all the column numbers that have 'Périmètre' in their name
        perimeter_columns = [i for i, x in enumerate(perimeter_row) if 'Périmètre' in str(x)]
        # Computing number of rows in table
        # Find the row number of the first table header ('ID')
        id_index = (full_df[0] == 'ID').idxmax()
        id_column = full_df.loc[id_index+1:, 0]
        # Count the number of rows in the ID column
        rows = id_column.count()
        # Compute number of columns in table
        try:
            # By finding the CUT column
            cols = full_df.iloc[HEADER_LINE, :].tolist().index('CUT')
        except ValueError:
            # ALl columns
            cols = full_df.shape[1]
        return {
            'modelName': modelName,
            'cols': cols,
            'rows': rows,
            'perimeterColumns': perimeter_columns
        }

    def get_nomenclature(elem):
        # filename to target (.csv format)
        nomenclature_name = elem['Détails de format'][elem['Détails de format'].index(':')+1:].strip()
        path_file = ''
        nomenclature_files = os.listdir(os.path.join("..", "nomenclature_parser", "out", "latest", "csv"))
        file_extension = ".csv"
        for filename in nomenclature_files:
            if filename[:-len(file_extension)] == nomenclature_name:
                path_file = os.path.join("..", "nomenclature_parser", "out", "latest", "csv", filename)

        if path_file != '':
            df_nomenclature = pd.read_csv(path_file, sep=",", keep_default_na=False, na_values=['_'], encoding="utf-8",
                                          dtype={'code': str})
            # Return codes with replaced commas and trimmed leading and trailing spaces for cleaner values
            L_ret = [code.replace(',', ' ').strip() for code in df_nomenclature["code"].values.tolist()]
        # ToDo: ajouter un bloc dans le elseif pour détecter des https:// et aller chercher les nomenclatures publiées en ligne (MOS/NOs par exemple)
        else:
            print(f"{Color.RED}ERROR: nomenclature {nomenclature_name} does not exist, could not load associated values.")
            print(f'Known nomenclatures are {nomenclature_files}')
            print("Check if some nomenclature files disappeared. If so, last run of nomenclature_parser.py likely failed.")
            exit(1)
        return L_ret

    params = get_params_from_sheet(sheet)
    # Schema name is in name = RC-EDA (or RS-EDA) for instance
    MODEL_NAME = params['modelName']  # CreateCase

    MODEL_TYPE = model_type
    WRAPPER_NAME = f"{MODEL_TYPE}Wrapper"  # createCaseWrapper
    NB_ROWS = params['rows']
    NB_COLS = params['cols']

    def is_custom_content():
        return MODEL_TYPE == "customContent"

    def is_allowing_additional_properties(name):
        return is_custom_content() or MODEL_TYPE == "DistributionElement" or name == "RC-DE"

    Path('out/' + name).mkdir(parents=True, exist_ok=True)

    # DATA COLLECTION AND CLEANING
    # Read CSV, skipping useless first and last lines
    df = pd.read_excel(filepath, sheet_name=sheet, skiprows=HEADER_LINE, nrows=NB_ROWS, converters={'ID': int})
    # Dropping useless columns
    df = df.iloc[:, :NB_COLS]
    # Column validation
    REQUIRED_COLUMNS = [
        *[f"Donnée (Niveau {i})" for i in range(1, DATA_DEPTH + 1)],
        'ID', 'Description', 'Exemples', 'Balise', 'Cardinalité', 'Objet', 'Format (ou type)', 'Détails de format',
        *([perimeter_filter] if perimeter_filter else [])
    ]
    if not (set(REQUIRED_COLUMNS) <= set(df.columns)):
        print(f"{Color.RED}ERROR: some key columns are missing:{Color.ORANGE}")
        print(set(REQUIRED_COLUMNS) - set(df.columns))
        print(f"Make sure all these columns are existing: {REQUIRED_COLUMNS} (headers on line {HEADER_LINE + 1}).{Color.END}")
        exit(1)

    # Keeping only the relevant perimeter rows
    if perimeter_filter:
        df = df[pd.notna(df[perimeter_filter])]
        # Replace 'Cardinalité' column values with the relevant perimeter column values (whenever the value is not 'X')
        df['Cardinalité'] = df.where(df[perimeter_filter] == 'X', df[perimeter_filter], axis=0)['Cardinalité']

    # Deleting perimeter columns. N.B: dropping a
    # column (obviously) reduces the length df.columns, so we iterate in reverse order
    for i in sorted(params['perimeterColumns'], reverse=True):
        df.drop(df.columns[i], axis=1, inplace=True)

    # Storing input data in a file to track versions
    # Before storing, we keep only useful columns and we also do not want to write line index to the file (index=False)
    INPUT_CSV_COLUMNS = REQUIRED_COLUMNS.copy()
    INPUT_CSV_COLUMNS.remove("ID")
    if perimeter_filter:
        INPUT_CSV_COLUMNS.remove(perimeter_filter)
    df[INPUT_CSV_COLUMNS].to_csv(f'out/{name}/{name}.input.csv', index=False)

    # Replacing comment cells (starting with '# ') with NaN in 'Donnée xx' columns
    df.iloc[:, 1:1 + DATA_DEPTH] = df.iloc[:, 1:1 + DATA_DEPTH].applymap(
        lambda x: pd.NA if str(x).startswith('# ') else x)


    def find_data_level(element):
        """Find the first data level of the element"""
        for i in range(1, DATA_DEPTH + 1):
            if str(element[f"Donnée (Niveau {i})"]) != 'nan':
                return i
        return 0

    def format_codeandlabel_properties(child, parent):
        code_file = parent['Détails de format']
        """ For 'Code', set code file name to the 'Détails de format' column, remove it from parent """
        if child['Balise'] == "code":
            child['Détails de format'] = code_file
            df.loc[parent.ID-1, 'Détails de format'] = 'nan'
        """Set the level of the child to be the level of the parent + 1"""
        if find_data_level(child) != find_data_level(parent)+1:
            child = shift_child_data_levels(child, parent)
        return copy.deepcopy(child)

    def shift_child_data_levels(child, parent):
        """Shift the data levels of the child by the shift value"""
        shift_difference = find_data_level(child) - find_data_level(parent)+1
        child_cpy = copy.deepcopy(child)
        for i in range(1, 1 + DATA_DEPTH):
            if 0 < i-shift_difference <= DATA_DEPTH:
                child_cpy[f"Donnée (Niveau {i})"] = child[f"Donnée (Niveau {i-shift_difference})"]
        return child_cpy

    global first_codeandlabel_properties
    first_codeandlabel_properties = []
    global first_codeandlabel_name
    first_codeandlabel_name = ""

    def regenerate_ids(df):
        """Regenerate the IDs of the dataframe"""
        for index, row in df.copy().iterrows():
            df.at[index, 'ID'] = index + 1

    # Iterate over df and insert two lines after each entry of type 'codeAndLabel'
    for index, row in df.copy().iterrows():
        if row['Format (ou type)'] == 'codeAndLabel':
            # We save the children of the codeAndLabel (2 next rows) if it's the first one we've seen
            if not first_codeandlabel_properties:
                for i in range(1, 3):
                    first_codeandlabel_properties.append(df.loc[index + i].to_dict())
            # Otherwise, we add the children of the first codeAndLabel to the current codeAndLabel
            else:
                for i in range(1, 3):
                    # we check the highest level of data of the current row and make sure to modify
                    # the level of data of each property to be equal to the row's level + 1
                    prop_cpy = first_codeandlabel_properties[i - 1].copy()
                    prop_cpy['ID'] = row['ID']+i/10
                    df.loc[index + i/10] = format_codeandlabel_properties(prop_cpy, row)

    df.sort_index(axis=0, ascending=True, inplace=True, kind='quicksort')
    df.reset_index(drop=True, inplace=True)
    regenerate_ids(df)


    # Adding a name column ('Balise' by default)
    df['name'] = df['Balise']

    # DATA ENRICHMENT
    # Get level in data hierarchy
    df['level_shift'] = -1
    for i in range(1, 1 + DATA_DEPTH):
        # Number of distinct values in 'level i' column before that item included
        df[f"level_{i}"] = df[f"Donnée (Niveau {i})"].notnull().cumsum()
        # Number of distinct values in 'level i' column before that item excluded
        df[f"previous_level_{i}"] = df[f"level_{i}"].shift(periods=1, fill_value=0)
        # If the 2 numbers are different => that item's active 'Donnée' column is i => update 'level_shift' column
        df["level_shift"] = df.apply(
            lambda row: i if row[f"level_{i}"] != row[f"previous_level_{i}"] else row['level_shift'], axis=1)


    # DATA VALIDATION
    HAS_ERROR = False

    def get_row_donnees_count(df):
        # Filter columns based on the specified prefix
        niveau_columns = [col for col in df.columns if col.startswith('Donnée (Niveau')]
        # Count non-null values in each row for the selected columns
        return df[niveau_columns].notnull().sum(axis=1)

    row_donnees_count = get_row_donnees_count(df)
    # - Lines with multiple 'Données'
    if not df[row_donnees_count > 1].empty:
        print(f"{Color.RED}ERROR: some rows have multiple 'Données' fields:{Color.ORANGE}")
        print(df[row_donnees_count > 1])
        print(f"Perhaps these lines contains uncorrectly formatted comments (should start with '# ').{Color.END}")
        HAS_ERROR = True
    # - Lines with no 'Données'
    if not df[row_donnees_count == 0].empty:
        print(f"{Color.RED}ERROR: some rows have no 'Données' field:{Color.ORANGE}")
        print(df[row_donnees_count == 0])
        print(f"Perhaps these lines should be deleted or contained only comments.{Color.END}")
        HAS_ERROR = True
    # - No name
    if not df[df['name'].isnull()].empty:
        print(f"{Color.RED}ERROR: some rows have no 'name' field:{Color.ORANGE}")
        print(df[df['name'].isnull()])
        print(f"Name is based on column 'Balise'.\n"
              f"Check that this column is correctly set up.{Color.END}")
        HAS_ERROR = True
    # - name with spaces
    test = df[df['name'].str.contains(' ')]
    if not df[df['name'].str.contains(' ')].empty:
        print(f"{Color.RED}ERROR: some rows have spaces in their 'name' field:{Color.ORANGE}")
        print(df[df['name'].str.contains(' ')])
        print(f"Name is based on column 'Balise'.\n"
              f"Check that this column is correctly set up.{Color.END}")
        HAS_ERROR = True
    # - objects with basic types
    basic_types = ['integer', 'number', 'string', 'datetime', 'date', 'boolean']
    objects_with_basic_type_df = df.loc[
        (df['Objet'] == 'X') &
        (df['Format (ou type)'].isin(basic_types))
        ]
    if not objects_with_basic_type_df.empty:
        print(f"{Color.RED}ERROR: some rows have objects with basic type:{Color.ORANGE}")
        print(objects_with_basic_type_df)
        print(
            f"Objects (column 'Objet' = X) can't have a basic type (column 'Format (ou type)') {basic_types}.{Color.END}")
        HAS_ERROR = True
    # - Failed to compute level_shift
    if not df[df['level_shift'] == -1].empty:
        print(f"{Color.RED}ERROR: level_shift couldn't be computed for some rows:{Color.ORANGE}")
        print(df[df['level_shift'] == -1])
        print(
            f"If these rows were not mentioned in a previous error, "
            f"the exact reason should be investigated by the dev team...{Color.END}")
        HAS_ERROR = True
    # - 'file' is not handled correctly by OpenAPI generator
    # Ref.: https://github.com/OpenAPITools/openapi-generator/pull/5817
    if (df['name'] == 'file').sum() > 0:
        print(f"{Color.RED}ERROR: 'file' can't be used as a balise name:{Color.ORANGE}")
        print(df[df['name'] == 'file'])
        print(f"Consider renaming this balise, adding a prefix, ...{Color.END}")
        HAS_ERROR = True
    if HAS_ERROR:
        exit(1)

    # Get IDs, parent and build data hierarchy tree structure
    def get_parent(row):
        parent = '1'
        for i in range(1, row['level_shift']):
            parent += '.' + str(row[f"level_{i}"])
        return parent

    def is_typed_object(row):
        """Is elem an object and using a specific reusable type?"""
        isObject = row['Objet'] == 'X'
        isTyped = str(row['Format (ou type)']) != 'nan'
        return isObject & isTyped

    def is_health_only(row):
        # TODO: Rework this method, it's no longer applicable as is, as we have many more perimeters than just
        #  '15-15' and '15-18'.
        # """Is elem only for 15-15?"""
        # isHealthOnly = row['15-18'] != 'X' and row['15-15'] == 'X'
        # return isHealthOnly
        return False

    def capitalizeFirstLetter(string):
        return string[0].upper() + string[1:]

    def get_true_type(row):
        """Get the type of elem (defaults to its name if there is no type specified)"""
        if row['Format (ou type)'] == 'codeAndLabel':
            return row['name']
        elif row["Objet"] != "X" or row['is_typed_object']:
            return row["Format (ou type)"]
        return row['name']

    def get_parent_type(row):
        if row['level_shift'] == 1:
            return MODEL_TYPE
        return df.loc[row['parent']]['true_type']

    def build_id(row):
        id = '1'
        for i in range(1, row['level_shift'] + 1):
            id += '.' + str(row[f"level_{i}"])
        return id

    def build_full_name(row):
        return row[f"Donnée (Niveau {row['level_shift']})"]

    df['is_typed_object'] = df.apply(is_typed_object, axis=1)
    df['is_health_only'] = df.apply(is_health_only, axis=1)
    df['true_type'] = df.apply(get_true_type, axis=1)
    if df.size > 0:
        df['id'] = df.apply(build_id, axis=1)
        df = df.set_index('id', drop=False)
        df['parent'] = df.apply(get_parent, axis=1)
        df['parent_type'] = df.apply(get_parent_type, axis=1)
        df['full_name'] = df.apply(build_full_name, axis=1)



    # Verify that cardinality is formatted correctly (e.g. '0..1', '1..1', '0..n', '1..n'; regex (\d+..(\d+|n)))
    def validate_cardinality_format(cardinality):
        if re.match(r'^\d+\.\.(\d+|n)$', cardinality) and validate_cardinality_values(cardinality):
            return True
        return False

    # Verify that cardinality value is valid (first number is lower than second number or second number is 'n')
    def validate_cardinality_values(cardinality):
        cardinality_values = cardinality.split('..')
        if cardinality_values[1] == 'n':
            return True
        if int(cardinality_values[0]) <= int(cardinality_values[1]):
            return True
        return False

    # Validate cardinality for each row of df
    errs = []
    for index, row in df.iterrows():
        if not validate_cardinality_format(row['Cardinalité']):
            errs.append(row)
    if errs:
        print(f"{Color.RED}ERROR: some rows have invalid cardinality values in sheet {sheet}, perimeter {perimeter_filter}:{Color.ORANGE}")
        for err in errs:
            print("Row ID: ", f"{str(int(err['ID'])):<5s}", "Name: ", f"{err['name']:<25s}", "Cardinality: ", err['Cardinalité'])
        print(f"{Color.RED}Cardinalities should be formatted as '0..1', '1..1', '0..n', '1..n'.{Color.END}")
        exit(1)


    # 2. Recursive data (children in their parent, to be explored like a tree)
    def get_element_with_its_children(previous_children, elem_id):
        elem = df.loc[elem_id].to_dict()
        if elem['id'] in previous_children:
            elem['children'] = previous_children[elem_id]
        return elem

    children = {}

    if not is_custom_content():
        for i in range(DATA_DEPTH, 0, -1):
            previous_children = children
            children_df = df[df['level_shift'] == i]
            children_ids = children_df.groupby('parent')['id'].apply(list)
            children = {
                parent_id: list(
                    map(lambda child_id: get_element_with_its_children(previous_children, child_id), children_ids))
                for (parent_id, children_ids) in children_ids.items()}
        rootObject = {
            'id': '1',
            'level_shift': 0,
            'name': MODEL_TYPE,
            'Objet': 'X',
            'children': children['1']
        }
    else:
        rootObject = {
            'id': '1',
            'level_shift': 0,
            'name': MODEL_TYPE,
            'Objet': 'X',
            'children': {}
        }

    # DATA USAGE
    def is_array(elem):
        """Is elem an array?"""
        cardinality = elem['Cardinalité']
        return not cardinality.endswith('1')

    def navigate_children_with_id(id_path):
        """
        Helper to navigate through children tree using item id.
        ex : '1.1.11' gives children['1'][0]['children'][11]['children']
        """
        # Removing wrapper
        current = children['1']
        id_path = id_path[2:]
        # Navigating following given id_path
        keys = list(map(int, id_path.split('.')))
        for key in keys:
            current = current[key - 1]['children']
        return current

    # Recursively build a json example
    def build_example(elem):
        if elem['Objet'] != 'X':
            if str(elem['Exemples']) == 'nan':
                return 'None'
            return elem['Exemples']
        elif 'children' not in elem:
            # ToDo: Could be object reference (so no children itself but first definition of its type has)
            # But it is not so easy to access the first reference here as nor the JSON Schema nor example are built yet
            # .id is not correct as the ID don't reset but are always increasing... computing some resetting id is necessary
            # type_first_definition = df[df['true_type'] == elem['true_type']].iloc[0].id
            # type_children = navigate_children_with_id(type_first_definition)
            return {}
        else:
            children = {}
            for child in elem['children']:
                if is_array(child):
                    children[child['name']] = [build_example(child)]
                else:
                    children[child['name']] = build_example(child)
            return children

    json_example = build_example(rootObject)
    with open(f'out/{name}/{name}.example.json', 'w', encoding='utf8') as outfile:
        json.dump(json_example, outfile, indent=4, ensure_ascii=False)

    # Go through data (list or tree) and use it to build the expected JSON schema
    json_schema = {
        '$schema': 'http://json-schema.org/draft-07/schema#',
        '$id': 'classpath:/json-schema/schema#',
        'x-id': f'{name}.schema.json#',  # required by JSV to find the schema file locally
        'version': version,
        'example': 'example.json#',
        'type': 'object',
        'title': MODEL_TYPE,
        'required': [],
        'properties': {},
        'definitions': {},
        'additionalProperties': is_allowing_additional_properties(name)
    }

    def has_format_details(elem, details):
        """Does elem have a format details starting with details?"""
        return str(elem['Détails de format']) != 'nan' and elem['Détails de format'].startswith(details)

    def type_matching(child):
        """Get the matching type for a given type name"""
        typeName = child['Format (ou type)']
        if typeName == 'date':
            return 'string', r'^\d{4}-\d{2}-\d{2}$', None
        elif typeName == 'datetime':
            return 'string', r'^\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}[\-+]\d{2}:\d{2}$', 'date-time'
        elif typeName == 'phoneNumber':
            return 'string', r'^tel:([#\+\*]|37000|00+)?[0-9]{2,15}$', None
        else:
            if has_format_details(child, FormatFlags.REGEX):
                return typeName, child['Détails de format'][child['Détails de format'].index(':')+1:].strip(), None
            else:
                return typeName, None, None

    def get_parent_example_path(parent):
        if parent['level_shift'] == 0:
            return json_schema['example']
        return json_schema['definitions'][parent['true_type']]['example']

    def child_not_already_required(child, definitions):
        """Check if the array of required values already contains the child we're attempting to add, return False if
        so"""
        return child['name'] not in definitions['required']

    def add_field_child_property(parent, child, definitions):
        """Update parent definitions (required and properties) by adding the child information for a field child"""
        if child['Cardinalité'].startswith('1') and child_not_already_required(child, definitions):
            definitions['required'].append(child['name'])
        typeName, pattern, format = type_matching(child)
        parentExamplePath = get_parent_example_path(parent)
        childDetails = {
            'type': typeName,
            'title': child['full_name'],
            'x-health-only': child['is_health_only'],
            'x-cols': 6,
            'example': parentExamplePath + '/' + child['name'] + ('/0' if is_array(child) else '')
        }
        if str(child['Description']) != 'nan':
            childDetails['description'] = child['Description']
        if pattern is not None:
            childDetails['pattern'] = pattern
        if format is not None:
            childDetails['format'] = format
        if has_format_details(child, FormatFlags.ENUM):
            childDetails['enum'] = child['Détails de format'][child['Détails de format'].index(':')+1:].strip().split(', ')
        # key word nomenclature trigger search over nomenclature folder for matching file
        if has_format_details(child, FormatFlags.NOMENCLATURE):
            childDetails['enum'] = get_nomenclature(child)
        properties = definitions['properties']
        if is_array(child):
            properties[child['name']] = {
                'type': 'array',
                'x-health-only': child['is_health_only'],
                'items': childDetails
            }
            if child['Cardinalité'][-1].isdigit():
                properties[child['name']]['maxItems'] = int(child['Cardinalité'][-1])
            if not child['Cardinalité'].startswith('0'):
                properties[child['name']]['minItems'] = int(child['Cardinalité'][0])
        else:
            properties[child['name']] = childDetails

    def add_object_child_definition(parent, child, definitions):
        """
        Update parent definitions (required and properties) by adding the child information for an object child
        Creates definitions for the child object if it does not exist yet.
        Special case: sourceMessage in rs-error message should allow additional properies (by default it is an
        empty object with no required or optional properties)
        """
        childOriginalTypeName = child['Format (ou type)']
        parentExamplePath = get_parent_example_path(parent)
        childTrueTypeName = child['true_type']
        if childTrueTypeName not in json_schema['definitions']:
            json_schema['definitions'][childTrueTypeName] = {
                'type': 'object',
                'title': child['full_name'],
                'x-display': 'expansion-panels',
                'x-health-only': child['is_health_only'],
                'required': [],
                'properties': {},
                'additionalProperties':  is_source_message(childTrueTypeName),
                'example': parentExamplePath + '/' + child['name'] + ('/0' if is_array(child) else '')
            }
        elif (childOriginalTypeName != "codeAndLabel"
              and 'children' in child):
            """If this is not the first occurrence of the object
            and it has children, then the model is incorrectly defined and we should throw an error and exit.
            We make an exception for codeAndLabel"""
            print(f"{Color.RED}ERROR: object '{childTrueTypeName}' is defined multiple times. ")
            print(f"Make sure the object is not defined multiple times in the model.{Color.END}")
            exit(1)

        """If this is the first codeAndLabel, we record its name, otherwise we copy the properties from the first 
        codeAndLabel to the current codeAndLabel"""
        if childOriginalTypeName == "codeAndLabel":
            global first_codeandlabel_name
            if first_codeandlabel_name == "":
                first_codeandlabel_name = childTrueTypeName
            else:
                json_schema['definitions'][childTrueTypeName]['properties'] = json_schema['definitions'][first_codeandlabel_name]['properties'].copy()

        if child['Cardinalité'].startswith('1'):
            definitions['required'].append(child['name'])
        properties = definitions['properties']
        if is_array(child):
            properties[child['name']] = {
                'type': 'array',
                'items': {
                    '$ref': '#/definitions/' + childTrueTypeName,
                }
            }
            if child['Cardinalité'][-1].isdigit():
                properties[child['name']]['maxItems'] = int(child['Cardinalité'][-1])
            if not child['Cardinalité'].startswith('0'):
                properties[child['name']]['minItems'] = int(child['Cardinalité'][0])
        else:
            properties[child['name']] = {
                '$ref': '#/definitions/' + childTrueTypeName,
            }

    def is_source_message(typeName):
        if typeName == "sourceMessage":
            return True
        return False

    def add_child(parent, child, definitions):
        """Update parent definitions by adding the child information"""
        if child['Objet'] != 'X':
            add_field_child_property(parent, child, definitions)
        else:
            add_object_child_definition(parent, child, definitions)

    def fill_object_definition(elem, root=False):
        """
        Update the definition (required and properties) of the type of elem by traversing its direct children
        if root, write it on the top level of the schema and not in definitions
        """
        if root:
            definition = json_schema
        else:
            if 'children' not in elem:
                """If element type is 'codeAndLabel' we check by 'name', else we check by 'Format (ou type)'"""
                if elem['Format (ou type)'] != 'codeAndLabel':
                    assert elem['Format (ou type)'] in json_schema['definitions'], \
                        (f"The type of the object '{elem['name']}' is not defined.\n"
                        f"Make sure the object is not empty")
                    return json_schema['definitions'][elem['Format (ou type)']]
                else:
                    assert elem['name'] in json_schema['definitions'], \
                        (f"The type of the object '{elem['name']}' is not defined.\n"
                        f"Make sure the object is not empty")
                    return json_schema['definitions'][elem['name']]
            typeName = elem['true_type']
            definition = json_schema['definitions'][typeName]
        # Fill the elem definitions based on the children
        for child in elem['children']:
            add_child(elem, child, definition)
        return definition

    def depth(x):
        """Get the depth of a dictionary"""
        if type(x) is dict and x:
            return 1 + max(depth(x[a]) for a in x)
        if type(x) is list and x:
            return 1 + max(depth(a) for a in x)
        return 0

    def get_examples_with_json_example(definition):
        """Collect example in the json_example if it is not too deep"""
        json_schema_path = "$" + definition['example'].split('#')[1]
        path_with_arrays = json_schema_path.replace('/0', '[0]')
        path = path_with_arrays.replace('/', '.')
        try:
            example = parse(path).find(json_example)[0].value
        except Exception as err:
            print(f"{Color.ORANGE}WARNING: '{err}' while getting example from JSON sample.{Color.ORANGE}")
            print(
                f"Path {path} of data '{definition['title']}' not found. Make sure parents and types are OK.{Color.END}")
            return
        if depth(example) < 2:
            definition['examples'] = [example]
        # Do it as well for all its children properties
        if 'properties' in definition:
            for prop in definition['properties']:
                propDetails = definition['properties'][prop]
                if 'items' in propDetails:
                    get_examples_with_json_example(get_definition_from_prop(propDetails['items']))
                else:
                    get_examples_with_json_example(get_definition_from_prop(propDetails))

    def get_definition_from_prop(prop):
        """Gives the definition of a property"""
        if '$ref' in prop:
            return json_schema['definitions'][prop['$ref'].split('/')[-1]]
        else:
            return prop

    def build_json_schema(elem):
        # BUILD JSON SCHEMA
        if elem['Objet'] == 'X':
            # Root element: write it on the top level of the schema and not in definitions
            # Lower objects: in definitions -> update the type in definitions by traversing the direct children
            fill_object_definition(elem, root=elem['level_shift'] == 0)

    def DFS(root, use_elem):
        use_elem(root)
        if 'children' in root:
            for child in root['children']:
                DFS(child, use_elem)

    print(f'{Color.BOLD}{Color.UNDERLINE}{Color.PURPLE}Generating JSON schema...{Color.END}')
    DFS(rootObject, build_json_schema)

    """Before dumping, we verify the json schema definitions to make sure no objects with no properties are defined.
    This verification is skipped for RS-ERROR schema"""
    if MODEL_TYPE != "error":
        empty_object_errors = []
        for key, definition in json_schema['definitions'].items():
            if not definition['properties']:
                empty_object_errors.append(f"{Color.RED}ERROR: object '{key}' is defined but has no properties.{Color.END}")
        if empty_object_errors:
            for error in empty_object_errors:
                print(error)
            print(f"{Color.RED}Make sure no empty objects are defined in the model.{Color.END}")
            exit(1)
    with open(f'out/{name}/{name}.schema.json', 'w', encoding='utf8') as outfile:
        json.dump(json_schema, outfile, indent=4, ensure_ascii=False)
    print('JSON schema generated.')

    # BUILD OpenAPI SCHEMA
    print(f'{Color.BOLD}{Color.UNDERLINE}{Color.PURPLE}Generating OpenAPI schema...{Color.END}')
    def json_schema_to_openapi_schema(json_schema):
        openapi_schema = {}
        definitions = json_schema['definitions']
        # Add root object to definitions so it is treated as a normal object
        # Dropping useless root infos
        root_definition = {key: json_schema[key] for key in json_schema if key not in [
            '$schema', 'definitions', 'version', 'id'
        ]}
        definitions = {
            **{MODEL_TYPE: root_definition},
            **definitions
        }
        # Simply collect all objects (and root properties)
        for elem_name, definition in definitions.items():
            get_examples_with_json_example(definition)
            openapi_schema[elem_name] = definition
        return openapi_schema

    openapi_components = json_schema_to_openapi_schema(json_schema)
    with open('RC-DE.openapi.yaml') as f:
        common_openapi_components = yaml.load(f, Loader=yaml.loader.SafeLoader)

    with open('template.openapi.yaml') as f:
        full_yaml = yaml.load(f, Loader=yaml.loader.SafeLoader)

        # Do not append wrapper_yaml if we're generating RC-DE schema
        if name != 'RC-DE':
            wrapper_yaml = {
                WRAPPER_NAME: {
                    "type": "object",
                    "required": [MODEL_TYPE],
                    "properties": {
                        MODEL_TYPE: {
                            "$ref": "#/components/schemas/" + MODEL_TYPE
                        }
                    }
                }
            }
        else:
            wrapper_yaml = {}

        full_yaml['components']['schemas'] = {
            **full_yaml['components']['schemas'],
            **wrapper_yaml,
            **openapi_components
        }

    with open(f'out/{name}/{name}.openapi.yaml', 'w', encoding='utf8') as file:
        documents = yaml.dump(full_yaml, sort_keys=False)
        documents = documents.replace('#/definitions/', "#/components/schemas/")
        file.write(documents)
    print('OpenAPI schema generated.')

    print(f'{Color.BOLD}{Color.UNDERLINE}{Color.PURPLE}Generating UML diagrams...{Color.END}')
    uml_generator.run(name, MODEL_TYPE, version=version)
    print('UML diagrams generated.')

    if not is_custom_content():
        named_df = df.copy().set_index(['parent_type', 'name']).fillna('')
    else:
        named_df = df.copy().set_index(['name']).fillna('')

    def get_excel_line(parent_type, name):
        try:
            # iloc[0] necessary even if there is only one line per name as it returns a DataFrame
            return named_df.loc[(parent_type, name)].iloc[0].to_dict()
        except AttributeError:
            return named_df.loc[(parent_type, name)].to_dict()

    def set_col_widths(table, widths):
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = docx.shared.Cm(width)

    def def_to_table(name, definition, title='', doc=None, style='Medium Shading 1 Accent 1'):
        if doc is None:
            doc = docx.Document()

        # Add title
        doc.add_heading(name, level=1)

        # Add paragraph
        # doc.add_paragraph('This table represents the fields and types defined in the JSON schema.')

        table = doc.add_table(rows=1, cols=6, style=style)

        # Add table header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Nom de balise'
        header_cells[1].text = 'Champ correspondant'
        header_cells[2].text = 'Format'
        header_cells[3].text = 'Cardinalité'
        header_cells[4].text = 'Description'
        header_cells[5].text = 'Exemple'

        # Iterate through each property in the JSON schema
        for field, properties in definition['properties'].items():
            field_data = get_excel_line(name, field)
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = field_data['full_name']
            typeInfo = field_data['true_type']
            if field_data['Objet'] == 'X':
                typeInfo = 'cf. type ' + typeInfo
            if field_data['Détails de format'] != '':
                typeInfo += "\n(" + field_data['Détails de format'] + ")"
            row_cells[2].text = typeInfo
            row_cells[3].text = field_data['Cardinalité']
            row_cells[4].text = field_data['Description']
            row_cells[5].text = str(field_data['Exemples'])

        # Specify the column widths (has to be on cells for Word) | Ref.: https://stackoverflow.com/a/43053996
        table.autofit = False
        table.allow_autofit = False
        column_widths = (3, 3.5, 2, 2.5, 8, 3)  # Widths in centimeters
        set_col_widths(table, column_widths)

        return doc

    print(f'{Color.BOLD}{Color.UNDERLINE}{Color.PURPLE}Generating Docx tables...{Color.END}')
    doc = docx.Document()
    # Set the page orientation to landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # Json Schema rootObject makes the object table
    def_to_table(MODEL_TYPE, json_schema, title=f"Objet {MODEL_NAME}", doc=doc)
    # Then all Json Schema definitions are types tables
    for elem_name, definition in json_schema['definitions'].items():
        def_to_table(elem_name, definition, title=f"Type {elem_name}", doc=doc)
    doc.save(f'out/{name}/{name}.schema.docx')

    print('Docx tables generated.')

    if RUN_DOCX_OUTPUT_EXAMPLES:
        # Build styles comparison
        for style in [
            # 'Table Normal',
            'Colorful Grid',
            'Colorful Grid Accent 1',
            'Colorful Grid Accent 2',
            'Colorful Grid Accent 3',
            'Colorful Grid Accent 4',
            'Colorful Grid Accent 5',
            'Colorful Grid Accent 6',
            'Colorful List',
            'Colorful List Accent 1',
            'Colorful List Accent 2',
            'Colorful List Accent 3',
            'Colorful List Accent 4',
            'Colorful List Accent 5',
            'Colorful List Accent 6',
            'Colorful Shading',
            'Colorful Shading Accent 1',
            'Colorful Shading Accent 2',
            'Colorful Shading Accent 3',
            'Colorful Shading Accent 4',
            'Colorful Shading Accent 5',
            'Colorful Shading Accent 6',
            'Dark List',
            'Dark List Accent 1',
            'Dark List Accent 2',
            'Dark List Accent 3',
            'Dark List Accent 4',
            'Dark List Accent 5',
            'Dark List Accent 6',
            'Light Grid',
            'Light Grid Accent 1',
            'Light Grid Accent 2',
            'Light Grid Accent 3',
            'Light Grid Accent 4',
            'Light Grid Accent 5',
            'Light Grid Accent 6',
            'Light List',
            'Light List Accent 1',
            'Light List Accent 2',
            'Light List Accent 3',
            'Light List Accent 4',
            'Light List Accent 5',
            'Light List Accent 6',
            'Light Shading',
            'Light Shading Accent 1',
            'Light Shading Accent 2',
            'Light Shading Accent 3',
            'Light Shading Accent 4',
            'Light Shading Accent 5',
            'Light Shading Accent 6',
            'Medium Grid 1',
            'Medium Grid 1 Accent 1',
            'Medium Grid 1 Accent 2',
            'Medium Grid 1 Accent 3',
            'Medium Grid 1 Accent 4',
            'Medium Grid 1 Accent 5',
            'Medium Grid 1 Accent 6',
            'Medium Grid 2',
            'Medium Grid 2 Accent 1',
            'Medium Grid 2 Accent 2',
            'Medium Grid 2 Accent 3',
            'Medium Grid 2 Accent 4',
            'Medium Grid 2 Accent 5',
            'Medium Grid 2 Accent 6',
            'Medium Grid 3',
            'Medium Grid 3 Accent 1',
            'Medium Grid 3 Accent 2',
            'Medium Grid 3 Accent 3',
            'Medium Grid 3 Accent 4',
            'Medium Grid 3 Accent 5',
            'Medium Grid 3 Accent 6',
            'Medium List 1',
            'Medium List 1 Accent 1',
            'Medium List 1 Accent 2',
            'Medium List 1 Accent 3',
            'Medium List 1 Accent 4',
            'Medium List 1 Accent 5',
            'Medium List 1 Accent 6',
            'Medium List 2',
            'Medium List 2 Accent 1',
            'Medium List 2 Accent 2',
            'Medium List 2 Accent 3',
            'Medium List 2 Accent 4',
            'Medium List 2 Accent 5',
            'Medium List 2 Accent 6',
            'Medium Shading 1',
            'Medium Shading 1 Accent 1',
            'Medium Shading 1 Accent 2',
            'Medium Shading 1 Accent 3',
            'Medium Shading 1 Accent 4',
            'Medium Shading 1 Accent 5',
            'Medium Shading 1 Accent 6',
            'Medium Shading 2',
            'Medium Shading 2 Accent 1',
            'Medium Shading 2 Accent 2',
            'Medium Shading 2 Accent 3',
            'Medium Shading 2 Accent 4',
            'Medium Shading 2 Accent 5',
            'Medium Shading 2 Accent 6',
        ]:
            if 'Accent 1' in style:
                def_to_table(MODEL_NAME, json_schema, style=style).save(f'docx-styles/schema-{style}.docx')
            else:
                def_to_table(MODEL_NAME, json_schema, style=style).save(f'docx-styles/others/schema-{style}.docx')

    # Add MODEL_TYPE to parsed schemas
    all_model_types.append(MODEL_TYPE)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        prog='Model Parser',
        description='Parses and converts the Excel model to the other needed formats',
    )
    parser.add_argument('-s', '--sheet', required=True, help='The Excel sheet to be parsed.')
    parser.add_argument('-n', '--name', required=True, help='The name to be given to the schema folder/file')
    parser.add_argument('-v', '--version', help='The version number to be used in model. Defaults to today.')
    parser.add_argument('-f', '--filter', default=False, help='If present, only 15-18 fields will be kept')
    args = parser.parse_args()

    run(args.sheet, args.name, args.version, args.filter)
